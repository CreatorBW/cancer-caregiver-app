import os
import re
import json
import sqlite3
import hashlib
from datetime import datetime, date
from typing import Any, Dict, List, Optional, Tuple

import fitz  # PyMuPDF
import docx
import pandas as pd
import streamlit as st
from dotenv import load_dotenv

try:
    from openai import OpenAI
except Exception:
    OpenAI = None


# ============================================================
# CONFIG
# ============================================================

load_dotenv()

APP_TITLE = "Cancer Caregiver Case Intelligence"
DB_PATH = "caregiver_case.db"

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")

client = OpenAI(api_key=OPENAI_API_KEY) if OpenAI and OPENAI_API_KEY else None


# ============================================================
# DATABASE
# ============================================================

def db():
    return sqlite3.connect(DB_PATH)


def execute(query: str, params: tuple = (), fetch: bool = False):
    conn = db()
    cur = conn.cursor()
    cur.execute(query, params)
    rows = cur.fetchall() if fetch else None
    conn.commit()
    conn.close()
    return rows


def init_db():
    conn = db()
    cur = conn.cursor()

    cur.execute("""
        CREATE TABLE IF NOT EXISTS app_meta (
            key TEXT PRIMARY KEY,
            value TEXT
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            file_name TEXT NOT NULL,
            file_hash TEXT UNIQUE NOT NULL,
            document_type TEXT,
            confidence TEXT,
            patient_name TEXT,
            document_date TEXT,
            primary_purpose TEXT,
            raw_text TEXT,
            classification_json TEXT,
            extraction_json TEXT,
            uploaded_at TEXT NOT NULL
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS case_profile (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            profile_key TEXT UNIQUE,
            profile_value TEXT,
            source_document TEXT,
            confidence TEXT,
            updated_at TEXT
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS conflicts (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            conflict_type TEXT,
            field_name TEXT,
            existing_value TEXT,
            new_value TEXT,
            source_document TEXT,
            status TEXT,
            created_at TEXT
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS diagnosis_records (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            diagnosis_date TEXT,
            primary_diagnosis TEXT,
            cancer_type TEXT,
            site TEXT,
            histology TEXT,
            stage TEXT,
            grade TEXT,
            metastatic_status TEXT,
            biomarkers_json TEXT,
            details TEXT,
            source_document TEXT,
            confidence TEXT,
            dedupe_key TEXT UNIQUE,
            created_at TEXT
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS timeline_events (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            event_date TEXT,
            event_type TEXT,
            title TEXT,
            details TEXT,
            source_document TEXT,
            confidence TEXT,
            dedupe_key TEXT UNIQUE,
            created_at TEXT
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS scan_events (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            scan_date TEXT,
            scan_type TEXT,
            body_region TEXT,
            findings TEXT,
            comparison TEXT,
            impression TEXT,
            disease_status TEXT,
            source_document TEXT,
            confidence TEXT,
            dedupe_key TEXT UNIQUE,
            created_at TEXT
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS lab_results (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            test_date TEXT,
            panel_name TEXT,
            test_name TEXT,
            value TEXT,
            unit TEXT,
            reference_range TEXT,
            flag TEXT,
            clinical_note TEXT,
            source_document TEXT,
            confidence TEXT,
            dedupe_key TEXT UNIQUE,
            created_at TEXT
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS medications (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            medicine_name TEXT,
            dose TEXT,
            route TEXT,
            frequency TEXT,
            duration TEXT,
            timing TEXT,
            purpose TEXT,
            start_date TEXT,
            end_date TEXT,
            special_instructions TEXT,
            source_document TEXT,
            confidence TEXT,
            dedupe_key TEXT UNIQUE,
            created_at TEXT
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS treatment_plans (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            plan_date TEXT,
            treatment_line TEXT,
            regimen_name TEXT,
            drugs_json TEXT,
            intent TEXT,
            cycle_info TEXT,
            next_cycle_date TEXT,
            required_tests_json TEXT,
            scan_plan TEXT,
            details TEXT,
            source_document TEXT,
            confidence TEXT,
            dedupe_key TEXT UNIQUE,
            created_at TEXT
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS side_effect_protocols (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            protocol_name TEXT,
            trigger_condition TEXT,
            steps_json TEXT,
            escalation_rule TEXT,
            source_document TEXT,
            confidence TEXT,
            dedupe_key TEXT UNIQUE,
            created_at TEXT
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS urgent_rules (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            risk_name TEXT,
            trigger_text TEXT,
            severity TEXT,
            recommended_action TEXT,
            source_document TEXT,
            confidence TEXT,
            dedupe_key TEXT UNIQUE,
            created_at TEXT
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS doctor_contacts (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT,
            role TEXT,
            phone TEXT,
            email TEXT,
            notes TEXT,
            source_document TEXT,
            confidence TEXT,
            dedupe_key TEXT UNIQUE,
            created_at TEXT
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS care_tasks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            task_date TEXT,
            task_type TEXT,
            title TEXT,
            details TEXT,
            status TEXT DEFAULT 'pending',
            source_document TEXT,
            confidence TEXT,
            dedupe_key TEXT UNIQUE,
            created_at TEXT
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS symptoms (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            symptom_date TEXT,
            symptom_name TEXT,
            severity INTEGER,
            temperature TEXT,
            notes TEXT,
            urgent_flags_json TEXT,
            created_at TEXT
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS caregiver_memory (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            memory_type TEXT,
            title TEXT,
            details TEXT,
            created_at TEXT,
            updated_at TEXT
        )
    """)

    conn.commit()
    conn.close()


def reset_database():
    conn = db()
    cur = conn.cursor()

    tables = [
        "documents",
        "case_profile",
        "conflicts",
        "diagnosis_records",
        "timeline_events",
        "scan_events",
        "lab_results",
        "medications",
        "treatment_plans",
        "side_effect_protocols",
        "urgent_rules",
        "doctor_contacts",
        "care_tasks",
        "symptoms",
        "caregiver_memory",
        "app_meta",
    ]

    for table in tables:
        cur.execute(f"DROP TABLE IF EXISTS {table}")

    conn.commit()
    conn.close()
    init_db()


# ============================================================
# UTILITIES
# ============================================================

def now_iso() -> str:
    return datetime.utcnow().isoformat()


def file_hash(file_bytes: bytes) -> str:
    return hashlib.sha256(file_bytes).hexdigest()


def dedupe_hash(*parts: Any) -> str:
    raw = "|".join([str(p or "").strip().lower() for p in parts])
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def safe_json_loads(raw: str) -> Dict[str, Any]:
    if not raw:
        return {}

    raw = raw.strip()

    if raw.startswith("```"):
        raw = raw.replace("```json", "").replace("```", "").strip()

    try:
        return json.loads(raw)
    except Exception:
        pass

    match = re.search(r"\{.*\}", raw, re.DOTALL)
    if match:
        return json.loads(match.group(0))

    raise ValueError("Could not parse JSON from model response.")


def as_list(value: Any) -> List[Any]:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    return [value]


def as_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def compact_text(value: Any, max_len: int = 5000) -> str:
    text = as_text(value)
    text = re.sub(r"\s+", " ", text).strip()
    return text[:max_len]


def parse_possible_date(value: Any) -> Optional[str]:
    if not value:
        return None

    text = str(value).strip()
    lower = text.lower()

    if lower in ["current", "latest", "recent", "now", "present"]:
        return None

    # Already ISO date
    if re.match(r"^\d{4}-\d{2}-\d{2}$", text):
        return text

    # YYYY-MM
    if re.match(r"^\d{4}-\d{2}$", text):
        return f"{text}-01"

    # Numeric date formats
    patterns = [
        ("%d.%m.%Y", r"\d{1,2}\.\d{1,2}\.\d{4}"),
        ("%d/%m/%Y", r"\d{1,2}/\d{1,2}/\d{4}"),
        ("%d-%m-%Y", r"\d{1,2}-\d{1,2}-\d{4}"),
        ("%d.%m.%y", r"\d{1,2}\.\d{1,2}\.\d{2}"),
        ("%d/%m/%y", r"\d{1,2}/\d{1,2}/\d{2}"),
        ("%d-%m-%y", r"\d{1,2}-\d{1,2}-\d{2}"),
    ]

    for fmt, pat in patterns:
        match = re.search(pat, text)
        if match:
            try:
                return datetime.strptime(match.group(0), fmt).date().isoformat()
            except Exception:
                pass

    # Month name + year
    month_map = {
        "january": "01", "jan": "01",
        "february": "02", "feb": "02",
        "march": "03", "mar": "03",
        "april": "04", "apr": "04",
        "may": "05",
        "june": "06", "jun": "06",
        "july": "07", "jul": "07",
        "august": "08", "aug": "08",
        "september": "09", "sep": "09", "sept": "09",
        "october": "10", "oct": "10",
        "november": "11", "nov": "11",
        "december": "12", "dec": "12",
    }

    # Handles: December 2021
    match = re.search(
        r"\b(january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec)\s*,?\s*(20\d{2}|19\d{2})\b",
        lower
    )
    if match:
        month = month_map.get(match.group(1))
        year = match.group(2)
        return f"{year}-{month}-01"

    # Handles: Jan 2023 - Feb 2024 / January 2023 to February 2024
    range_match = re.search(
        r"\b(january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec)\s*,?\s*(20\d{2}|19\d{2})\s*(?:-|to|until|till)",
        lower
    )
    if range_match:
        month = month_map.get(range_match.group(1))
        year = range_match.group(2)
        return f"{year}-{month}-01"

    # Handles: March/April 2024
    slash_month_match = re.search(
        r"\b(january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec)\s*/\s*(january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec)\s*,?\s*(20\d{2}|19\d{2})",
        lower
    )
    if slash_month_match:
        month = month_map.get(slash_month_match.group(1))
        year = slash_month_match.group(3)
        return f"{year}-{month}-01"

    return None


def timeline_bucket(raw_date: Any, title: Any, details: Any) -> str:
    text = f"{raw_date or ''} {title or ''} {details or ''}".lower()

    if any(x in text for x in ["current", "latest", "present status", "current status"]):
        return "Current / Latest"

    parsed = parse_possible_date(raw_date)
    if parsed:
        return "Chronological"

    return "Date Unclear"


def timeline_sort_date(raw_date: Any) -> str:
    parsed = parse_possible_date(raw_date)
    if parsed:
        return parsed

    text = str(raw_date or "").lower()

    if text in ["current", "latest", "recent", "now", "present"]:
        return "9999-12-31"

    return "9999-01-01"


def timeline_display_date(raw_date: Any) -> str:
    if raw_date in [None, ""]:
        return "Date unclear"
    return str(raw_date)


def make_timeline_display_df(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return df

    out = df.copy()

    out["Bucket"] = out.apply(
        lambda r: timeline_bucket(r.get("Date"), r.get("Title"), r.get("Details")),
        axis=1
    )

    out["Sort Date"] = out["Date"].apply(timeline_sort_date)
    out["Display Date"] = out["Date"].apply(timeline_display_date)

    # Collapse obvious duplicate display rows while keeping first source.
    # This prevents repeated same historical facts from multiple summaries dominating the timeline.
    out["Dedupe View Key"] = (
        out["Sort Date"].astype(str).str.lower().str.strip()
        + "|"
        + out["Type"].astype(str).str.lower().str.strip()
        + "|"
        + out["Title"].astype(str).str.lower().str.replace(r"\s+", " ", regex=True).str.strip()
    )

    out = out.drop_duplicates(subset=["Dedupe View Key"], keep="first")

    bucket_order = {
        "Current / Latest": 0,
        "Chronological": 1,
        "Date Unclear": 2,
    }

    out["Bucket Order"] = out["Bucket"].map(bucket_order).fillna(9)

    out = out.sort_values(
        by=["Bucket Order", "Sort Date", "Type", "Title"],
        ascending=[True, True, True, True]
    )

    return out.drop(columns=["Dedupe View Key", "Bucket Order"])

def insert_ignore(query: str, params: tuple):
    try:
        execute(query, params)
    except sqlite3.IntegrityError:
        pass


# ============================================================
# FILE TEXT EXTRACTION
# ============================================================

def extract_pdf_text(file_bytes: bytes) -> str:
    parts = []
    pdf = fitz.open(stream=file_bytes, filetype="pdf")

    for page_num, page in enumerate(pdf, start=1):
        text = page.get_text("text")
        if text and text.strip():
            parts.append(f"\n\n--- PAGE {page_num} ---\n{text.strip()}")

    return "\n".join(parts).strip()


def extract_docx_text(file_bytes: bytes) -> str:
    temp_path = "_temp_upload.docx"

    with open(temp_path, "wb") as f:
        f.write(file_bytes)

    document = docx.Document(temp_path)
    parts = []

    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    try:
        os.remove(temp_path)
    except Exception:
        pass

    return "\n".join(parts).strip()


def extract_txt_text(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="ignore")


def extract_text_from_file(file_name: str, file_bytes: bytes) -> str:
    lower = file_name.lower()

    if lower.endswith(".pdf"):
        return extract_pdf_text(file_bytes)

    if lower.endswith(".docx"):
        return extract_docx_text(file_bytes)

    if lower.endswith(".txt"):
        return extract_txt_text(file_bytes)

    raise ValueError("Unsupported file type. Upload PDF, DOCX, or TXT.")


# ============================================================
# AI PROMPTS
# ============================================================

CLASSIFICATION_SYSTEM = """
You classify uploaded medical documents for a cancer caregiver case-management application.

Return only valid JSON.

Do not diagnose.
Do not invent information.
Use the document text and file name only.
"""

UNIVERSAL_SCHEMA = """
{
  "document_classification": {
    "document_type": "discharge_summary | pet_ct_report | ct_report | mri_report | ultrasound_report | xray_report | lab_report | prescription | biopsy_report | histopathology_report | ihc_report | molecular_report | doctor_note | treatment_summary | bill_or_admin | unknown",
    "confidence": "high | medium | low",
    "patient_name": null,
    "document_date": null,
    "primary_purpose": null,
    "should_update": {
      "case_profile": false,
      "diagnosis_records": false,
      "treatment_plans": false,
      "timeline_events": false,
      "scan_events": false,
      "lab_results": false,
      "medications": false,
      "side_effect_protocols": false,
      "urgent_rules": false,
      "doctor_contacts": false,
      "care_tasks": false
    },
    "reason": null
  },
  "case_profile": {
    "patient_name": null,
    "age": null,
    "sex": null,
    "uhid_or_patient_id": null,
    "hospital": null,
    "department": null,
    "consultant": null,
    "allergies": [],
    "height": null,
    "weight": null,
    "bsa": null,
    "past_medical_history": null,
    "past_surgical_history": null
  },
  "diagnosis_records": [
    {
      "diagnosis_date": null,
      "primary_diagnosis": null,
      "cancer_type": null,
      "site": null,
      "histology": null,
      "stage": null,
      "grade": null,
      "metastatic_status": null,
      "biomarkers": [],
      "details": null,
      "confidence": "high | medium | low"
    }
  ],
  "treatment_plans": [
    {
      "plan_date": null,
      "treatment_line": null,
      "regimen_name": null,
      "drugs": [],
      "intent": null,
      "cycle_info": null,
      "next_cycle_date": null,
      "required_tests": [],
      "scan_plan": null,
      "details": null,
      "confidence": "high | medium | low"
    }
  ],
  "timeline_events": [
    {
      "event_date": null,
      "event_type": "diagnosis | treatment | surgery | radiation | chemotherapy | immunotherapy | targeted_therapy | progression | response | stable_disease | scan | lab | admission | discharge | follow_up | medication | procedure | symptom | other",
      "title": null,
      "details": null,
      "confidence": "high | medium | low"
    }
  ],
  "scan_events": [
    {
      "scan_date": null,
      "scan_type": null,
      "body_region": null,
      "findings": null,
      "comparison": null,
      "impression": null,
      "disease_status": "complete_response | partial_response | stable_disease | progressive_disease | no_active_disease | mixed_response | unclear",
      "confidence": "high | medium | low"
    }
  ],
  "lab_results": [
    {
      "test_date": null,
      "panel_name": null,
      "test_name": null,
      "value": null,
      "unit": null,
      "reference_range": null,
      "flag": "high | low | normal | abnormal | critical | unknown",
      "clinical_note": null,
      "confidence": "high | medium | low"
    }
  ],
  "medications": [
    {
      "medicine_name": null,
      "dose": null,
      "route": null,
      "frequency": null,
      "duration": null,
      "timing": null,
      "purpose": null,
      "start_date": null,
      "end_date": null,
      "special_instructions": null,
      "confidence": "high | medium | low"
    }
  ],
  "side_effect_protocols": [
    {
      "protocol_name": null,
      "trigger_condition": null,
      "steps": [],
      "escalation_rule": null,
      "confidence": "high | medium | low"
    }
  ],
  "urgent_rules": [
    {
      "risk_name": null,
      "trigger_text": null,
      "severity": "emergency | urgent | monitor",
      "recommended_action": null,
      "confidence": "high | medium | low"
    }
  ],
  "doctor_contacts": [
    {
      "name": null,
      "role": null,
      "phone": null,
      "email": null,
      "notes": null,
      "confidence": "high | medium | low"
    }
  ],
  "care_tasks": [
    {
      "task_date": null,
      "task_type": "medicine | injection | test | report_share | appointment | scan | monitoring | call_doctor | other",
      "title": null,
      "details": null,
      "confidence": "high | medium | low"
    }
  ],
  "caregiver_explanation": {
    "simple_summary": null,
    "what_changed": null,
    "what_to_do_next": [],
    "questions_for_doctor": [],
    "missing_or_unclear_items": []
  }
}
"""


def classify_document_ai(report_text: str, file_name: str) -> Dict[str, Any]:
    if not client:
        return classify_document_fallback(report_text, file_name)

    prompt = f"""
Classify this medical document.

Return ONLY JSON:

{{
  "document_type": "discharge_summary | pet_ct_report | ct_report | mri_report | ultrasound_report | xray_report | lab_report | prescription | biopsy_report | histopathology_report | ihc_report | molecular_report | doctor_note | treatment_summary | bill_or_admin | unknown",
  "confidence": "high | medium | low",
  "patient_name": null,
  "document_date": null,
  "primary_purpose": null,
  "should_update": {{
    "case_profile": false,
    "diagnosis_records": false,
    "treatment_plans": false,
    "timeline_events": false,
    "scan_events": false,
    "lab_results": false,
    "medications": false,
    "side_effect_protocols": false,
    "urgent_rules": false,
    "doctor_contacts": false,
    "care_tasks": false
  }},
  "reason": null
}}

File name:
{file_name}

Document text:
\"\"\"
{report_text[:18000]}
\"\"\"
"""

    response = client.chat.completions.create(
        model=OPENAI_MODEL,
        temperature=0,
        response_format={"type": "json_object"},
        messages=[
            {"role": "system", "content": CLASSIFICATION_SYSTEM},
            {"role": "user", "content": prompt}
        ],
    )

    result = safe_json_loads(response.choices[0].message.content)
    return normalize_classification(result)


def normalize_classification(c: Dict[str, Any]) -> Dict[str, Any]:
    doc_type = c.get("document_type") or "unknown"

    default_updates = update_map_for_doc_type(doc_type)

    should_update = c.get("should_update")
    if not isinstance(should_update, dict):
        should_update = default_updates
    else:
        merged = default_updates.copy()
        merged.update(should_update)
        should_update = merged

    return {
        "document_type": doc_type,
        "confidence": c.get("confidence") or "low",
        "patient_name": c.get("patient_name"),
        "document_date": parse_possible_date(c.get("document_date")) or c.get("document_date"),
        "primary_purpose": c.get("primary_purpose"),
        "should_update": should_update,
        "reason": c.get("reason"),
    }


def update_map_for_doc_type(doc_type: str) -> Dict[str, bool]:
    base = {
        "case_profile": False,
        "diagnosis_records": False,
        "treatment_plans": False,
        "timeline_events": False,
        "scan_events": False,
        "lab_results": False,
        "medications": False,
        "side_effect_protocols": False,
        "urgent_rules": False,
        "doctor_contacts": False,
        "care_tasks": False,
    }

    if doc_type == "discharge_summary":
        base.update({
            "case_profile": True,
            "diagnosis_records": True,
            "treatment_plans": True,
            "timeline_events": True,
            "scan_events": True,
            "lab_results": False,
            "medications": True,
            "side_effect_protocols": True,
            "urgent_rules": True,
            "doctor_contacts": True,
            "care_tasks": True,
        })

    elif doc_type in ["pet_ct_report", "ct_report", "mri_report", "ultrasound_report", "xray_report"]:
        base.update({
            "case_profile": True,
            "diagnosis_records": False,
            "treatment_plans": False,
            "timeline_events": True,
            "scan_events": True,
            "lab_results": False,
            "medications": False,
            "side_effect_protocols": False,
            "urgent_rules": False,
            "doctor_contacts": False,
            "care_tasks": False,
        })

    elif doc_type == "lab_report":
        base.update({
            "case_profile": True,
            "lab_results": True,
            "timeline_events": True,
            "care_tasks": False,
        })

    elif doc_type == "prescription":
        base.update({
            "case_profile": True,
            "medications": True,
            "side_effect_protocols": True,
            "urgent_rules": True,
            "care_tasks": True,
            "timeline_events": True,
        })

    elif doc_type in ["biopsy_report", "histopathology_report", "ihc_report", "molecular_report"]:
        base.update({
            "case_profile": True,
            "diagnosis_records": True,
            "timeline_events": True,
        })

    elif doc_type in ["doctor_note", "treatment_summary"]:
        base.update({
            "case_profile": True,
            "diagnosis_records": True,
            "treatment_plans": True,
            "timeline_events": True,
            "medications": True,
            "side_effect_protocols": True,
            "urgent_rules": True,
            "doctor_contacts": True,
            "care_tasks": True,
        })

    elif doc_type == "bill_or_admin":
        base.update({
            "case_profile": False,
            "diagnosis_records": False,
            "treatment_plans": False,
            "timeline_events": False,
            "scan_events": False,
            "lab_results": False,
            "medications": False,
            "side_effect_protocols": False,
            "urgent_rules": False,
            "doctor_contacts": False,
            "care_tasks": False,
        })

    else:
        base.update({
            "case_profile": True,
            "timeline_events": True,
        })

    return base


def classify_document_fallback(report_text: str, file_name: str) -> Dict[str, Any]:
    text = f"{file_name}\n{report_text}".lower()
    doc_type = "unknown"

    if "discharge summary" in text or ("date of admission" in text and "date of discharge" in text):
        doc_type = "discharge_summary"
    elif "pet ct" in text or "pet-ct" in text or "fdg" in text:
        doc_type = "pet_ct_report"
    elif "cect" in text or "ct scan" in text:
        doc_type = "ct_report"
    elif "mri" in text:
        doc_type = "mri_report"
    elif any(x in text for x in ["hemoglobin", "haemoglobin", "wbc", "platelet", "creatinine", "bilirubin", "sgot", "sgpt", "cbc", "lft", "kft"]):
        doc_type = "lab_report"
    elif any(x in text for x in ["rx", "prescription", "tab.", "cap.", "inj.", "syp."]):
        doc_type = "prescription"
    elif any(x in text for x in ["biopsy", "histopathology", "hpe"]):
        doc_type = "histopathology_report"
    elif any(x in text for x in ["ihc", "her2", "her-2", "pd-l1", "pdl1", "msi", "mmr"]):
        doc_type = "ihc_report"
    elif any(x in text for x in ["invoice", "bill", "receipt", "amount payable", "tax invoice"]):
        doc_type = "bill_or_admin"

    return normalize_classification({
        "document_type": doc_type,
        "confidence": "medium",
        "patient_name": None,
        "document_date": None,
        "primary_purpose": f"Rule-based classification as {doc_type}",
        "should_update": update_map_for_doc_type(doc_type),
        "reason": "Fallback keyword classification"
    })


EXTRACTION_SYSTEM = """
You are a medical document structuring engine for a cancer caregiver application.

Your job:
- Extract facts explicitly present in the uploaded document.
- Use the document classification to decide what sections to populate.
- Do not fill irrelevant sections just to be complete.
- Do not diagnose.
- Do not prescribe.
- Do not suggest medicine changes.
- Do not invent dates, doses, scan results, or treatment intent.
- Preserve uncertainty.
- If OCR is noisy, extract only what is reasonably clear and mark confidence lower.
- Return only valid JSON matching the requested schema.
"""


def extraction_focus_for_type(doc_type: str) -> str:
    if doc_type == "discharge_summary":
        return """
Focus on:
- patient identity and hospital details
- diagnosis and cancer status
- admission/discharge facts
- treatment given during admission
- treatment history/timeline
- current treatment plan
- next cycle date
- tests required before next cycle
- discharge medicines
- side-effect protocols
- emergency red flags
- doctor/coordinator contacts
- care tasks
- scan history if included
"""

    if doc_type in ["pet_ct_report", "ct_report", "mri_report", "ultrasound_report", "xray_report"]:
        return """
Focus on:
- scan date
- scan type
- body region
- key findings
- comparison with previous scan
- impression
- disease status: progression, stable disease, response, no active disease, mixed response, unclear
- timeline event for this scan
- doctor questions related to the scan
Do not create medication schedules unless medicines are explicitly part of the scan report.
"""

    if doc_type == "lab_report":
        return """
Focus on:
- test date
- panel name
- each lab parameter, value, unit, reference range, high/low/normal/critical flag if present
- clinical notes only if written in the report
- timeline event for lab report
Do not infer chemo fitness unless the report explicitly states it.
"""

    if doc_type == "prescription":
        return """
Focus on:
- medicine name
- dose
- route
- frequency
- timing
- duration
- start/end date if written
- purpose if explicitly clear from context
- special instructions
- side-effect protocols
- care tasks
- urgent instructions if written
Do not alter or simplify doses.
"""

    if doc_type in ["biopsy_report", "histopathology_report", "ihc_report", "molecular_report"]:
        return """
Focus on:
- diagnosis confirmation
- sample/site
- histology
- grade
- markers/biomarkers
- HER2/PD-L1/MSI/MMR/mutations if present
- interpretation only as written
- timeline event
Do not recommend treatment based on markers.
"""

    if doc_type in ["doctor_note", "treatment_summary"]:
        return """
Focus on:
- doctor instructions
- treatment plan
- next steps
- medicines if written
- tests/scans advised
- follow-up date
- warning signs if written
- tasks and timeline events
"""

    if doc_type == "bill_or_admin":
        return """
Focus on:
- classify as admin.
- Do not update clinical timeline unless there is explicit clinical information.
"""

    return """
Focus on extracting only clear facts. Populate only relevant sections. Mark unclear items.
"""


def ai_extract_document(report_text: str, file_name: str, classification: Dict[str, Any]) -> Dict[str, Any]:
    if not client:
        return fallback_extract_document(report_text, file_name, classification)

    doc_type = classification.get("document_type", "unknown")

    prompt = f"""
You are processing one uploaded medical document for a cancer caregiver case system.

Document classification:
{json.dumps(classification, ensure_ascii=False, indent=2)}

Extraction focus:
{extraction_focus_for_type(doc_type)}

Return ONLY valid JSON matching this universal schema:

{UNIVERSAL_SCHEMA}

Rules:
- Do not populate irrelevant sections.
- If this is a lab report, focus on lab_results.
- If this is a scan report, focus on scan_events.
- If this is a prescription, focus on medications and protocols.
- If this is a discharge summary, extract all relevant sections.
- If a date cannot be normalized to YYYY-MM-DD, keep the original text.
- Do not hardcode or assume anything.
- Use null or empty arrays when absent.
- Every extracted item should be traceable to the uploaded document.
- Confidence should reflect clarity of OCR/source text.

File name:
{file_name}

Document text:
\"\"\"
{report_text[:65000]}
\"\"\"
"""

    response = client.chat.completions.create(
        model=OPENAI_MODEL,
        temperature=0.05,
        response_format={"type": "json_object"},
        messages=[
            {"role": "system", "content": EXTRACTION_SYSTEM},
            {"role": "user", "content": prompt}
        ],
    )

    result = safe_json_loads(response.choices[0].message.content)
    return normalize_extraction(result, classification)


def normalize_extraction(data: Dict[str, Any], classification: Dict[str, Any]) -> Dict[str, Any]:
    data.setdefault("document_classification", classification)
    data["document_classification"] = normalize_classification(data.get("document_classification") or classification)

    for key in [
        "diagnosis_records",
        "treatment_plans",
        "timeline_events",
        "scan_events",
        "lab_results",
        "medications",
        "side_effect_protocols",
        "urgent_rules",
        "doctor_contacts",
        "care_tasks",
    ]:
        data[key] = as_list(data.get(key))

    if not isinstance(data.get("case_profile"), dict):
        data["case_profile"] = {}

    if not isinstance(data.get("caregiver_explanation"), dict):
        data["caregiver_explanation"] = {
            "simple_summary": None,
            "what_changed": None,
            "what_to_do_next": [],
            "questions_for_doctor": [],
            "missing_or_unclear_items": []
        }

    return data


# ============================================================
# FALLBACK EXTRACTION
# ============================================================

def fallback_extract_document(report_text: str, file_name: str, classification: Dict[str, Any]) -> Dict[str, Any]:
    doc_type = classification.get("document_type", "unknown")
    text = report_text

    extraction = {
        "document_classification": classification,
        "case_profile": fallback_case_profile(text),
        "diagnosis_records": [],
        "treatment_plans": [],
        "timeline_events": [],
        "scan_events": [],
        "lab_results": [],
        "medications": [],
        "side_effect_protocols": [],
        "urgent_rules": [],
        "doctor_contacts": [],
        "care_tasks": [],
        "caregiver_explanation": {
            "simple_summary": "Basic rule-based extraction completed because AI is unavailable.",
            "what_changed": None,
            "what_to_do_next": [],
            "questions_for_doctor": [],
            "missing_or_unclear_items": ["AI extraction was unavailable. Verify extracted data manually."]
        }
    }

    if doc_type == "discharge_summary":
        extraction["diagnosis_records"] = fallback_diagnosis(text)
        extraction["treatment_plans"] = fallback_treatment_plan(text)
        extraction["timeline_events"] = fallback_timeline(text)
        extraction["scan_events"] = fallback_scans(text)
        extraction["medications"] = fallback_medications(text)
        extraction["side_effect_protocols"] = fallback_protocols(text)
        extraction["urgent_rules"] = fallback_urgent_rules(text)
        extraction["doctor_contacts"] = fallback_contacts(text)
        extraction["care_tasks"] = fallback_tasks(text)

    elif doc_type in ["pet_ct_report", "ct_report", "mri_report", "ultrasound_report", "xray_report"]:
        extraction["scan_events"] = fallback_scans(text)
        extraction["timeline_events"] = fallback_timeline(text)

    elif doc_type == "lab_report":
        extraction["lab_results"] = fallback_labs(text)
        extraction["timeline_events"] = fallback_timeline(text)

    elif doc_type == "prescription":
        extraction["medications"] = fallback_medications(text)
        extraction["side_effect_protocols"] = fallback_protocols(text)
        extraction["urgent_rules"] = fallback_urgent_rules(text)
        extraction["care_tasks"] = fallback_tasks(text)

    elif doc_type in ["biopsy_report", "histopathology_report", "ihc_report", "molecular_report"]:
        extraction["diagnosis_records"] = fallback_diagnosis(text)
        extraction["timeline_events"] = fallback_timeline(text)

    else:
        extraction["timeline_events"] = fallback_timeline(text)

    return normalize_extraction(extraction, classification)


def regex_first(text: str, patterns: List[str]) -> Optional[str]:
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return compact_text(match.group(1), 400)
    return None


def fallback_case_profile(text: str) -> Dict[str, Any]:
    return {
        "patient_name": regex_first(text, [r"Patient'?s Name:\s*([^\n]+)", r"Name:\s*([^\n]+)"]),
        "age": regex_first(text, [r"Age:\s*([^\n]+)"]),
        "sex": regex_first(text, [r"Sex:\s*([^\n]+)", r"Gender:\s*([^\n]+)"]),
        "uhid_or_patient_id": regex_first(text, [r"UHID\s*No[:.]?\s*([^\n]+)", r"Patient\s*ID[:.]?\s*([^\n]+)"]),
        "hospital": regex_first(text, [r"(Fortis[^\n]+)", r"Hospital[:.]?\s*([^\n]+)"]),
        "department": regex_first(text, [r"Department\s*of\s*([^\n]+)", r"Department[:.]?\s*([^\n]+)"]),
        "consultant": regex_first(text, [r"Consultant[:.]?\s*([^\n]+)", r"Doctor[:.]?\s*([^\n]+)"]),
        "allergies": extract_allergies(text),
        "height": regex_first(text, [r"Height[:.]?\s*([^\n]+?)(?:Weight|$)"]),
        "weight": regex_first(text, [r"Weight[:.]?\s*([^\n]+?)(?:BSA|$)"]),
        "bsa": regex_first(text, [r"BSA[:.]?\s*([^\n]+)"]),
        "past_medical_history": regex_first(text, [r"Past Medical History[:.]?\s*-?\s*([^\n]+)"]),
        "past_surgical_history": regex_first(text, [r"Past Surgical History[:.]?\s*-?\s*([^\n]+)"]),
    }


def extract_allergies(text: str) -> List[str]:
    allergies = []
    lines = text.splitlines()

    for line in lines:
        lower = line.lower()
        if "allerg" in lower:
            clean = compact_text(line, 300)
            if clean:
                allergies.append(clean)

    return allergies[:10]


def fallback_diagnosis(text: str) -> List[Dict[str, Any]]:
    diagnosis = regex_first(text, [r"Diagnosis[:.]?\s*([^\n]+)"])
    stage = regex_first(text, [r"Stage[:.]?\s*([^\n]+)"])
    lower = text.lower()

    if not diagnosis and not any(x in lower for x in ["carcinoma", "cancer", "adenocarcinoma", "lymphoma", "sarcoma", "malignancy"]):
        return []

    return [{
        "diagnosis_date": None,
        "primary_diagnosis": diagnosis,
        "cancer_type": keyword_or_none(lower, ["adenocarcinoma", "carcinoma", "lymphoma", "sarcoma"]),
        "site": None,
        "histology": keyword_or_none(lower, ["adenocarcinoma", "squamous", "neuroendocrine"]),
        "stage": stage,
        "grade": regex_first(text, [r"Grade[:.]?\s*([^\n]+)"]),
        "metastatic_status": "metastatic" if "metastatic" in lower else None,
        "biomarkers": extract_biomarkers(text),
        "details": diagnosis or "Diagnosis-related terms found in document.",
        "confidence": "low"
    }]


def keyword_or_none(text: str, keywords: List[str]) -> Optional[str]:
    for k in keywords:
        if k in text:
            return k
    return None


def extract_biomarkers(text: str) -> List[str]:
    markers = []
    terms = ["HER2", "HER-2", "Her 2", "PD-L1", "PDL1", "MSI", "MMR", "EGFR", "ALK", "ROS1", "BRAF", "KRAS", "NRAS", "NTRK"]
    for term in terms:
        if term.lower() in text.lower():
            markers.append(term)
    return sorted(list(set(markers)))


def fallback_treatment_plan(text: str) -> List[Dict[str, Any]]:
    lower = text.lower()
    if not any(x in lower for x in ["chemotherapy", "radiotherapy", "radiation", "immunotherapy", "trastuzumab", "carboplatin", "cycle", "regimen", "treatment"]):
        return []

    drugs = []
    drug_candidates = [
        "Trastuzumab", "Carboplatin", "Cisplatin", "Oxaliplatin", "Paclitaxel", "Docetaxel",
        "Capecitabine", "Xeloda", "Herceptin", "Enhertu", "TDM1", "Pembrolizumab",
        "Nivolumab", "Bevacizumab", "FLOT", "5-FU"
    ]
    for d in drug_candidates:
        if d.lower() in lower:
            drugs.append(d)

    return [{
        "plan_date": None,
        "treatment_line": regex_first(text, [r"Plan[:.]?\s*([^\n]+)"]),
        "regimen_name": " + ".join(drugs) if drugs else None,
        "drugs": drugs,
        "intent": None,
        "cycle_info": regex_first(text, [r"Cycle[:.]?\s*([^\n]+)"]),
        "next_cycle_date": regex_first(text, [r"due on\s*([0-9./-]+)", r"next cycle[:.]?\s*([^\n]+)"]),
        "required_tests": extract_required_tests(text),
        "scan_plan": regex_first(text, [r"(PET\s*CT\s*Scan\s*to\s*be\s*done[^\n]+)", r"(scan\s*to\s*be\s*done[^\n]+)"]),
        "details": compact_text(regex_first(text, [r"Treatment Given[:.]?\s*([^\n]+)", r"Plan[:.]?\s*([^\n]+)"]), 1000),
        "confidence": "low"
    }]


def extract_required_tests(text: str) -> List[str]:
    tests = []
    for t in ["CBC", "LFT", "KFT", "RFT", "Creatinine", "Liver Function", "Kidney Function"]:
        if t.lower() in text.lower():
            tests.append(t)
    return tests


def fallback_timeline(text: str) -> List[Dict[str, Any]]:
    events = []
    lines = [compact_text(x, 600) for x in text.splitlines() if len(x.strip()) > 5]

    for line in lines:
        lower = line.lower()
        has_date = bool(re.search(r"\d{1,2}[./-]\d{1,2}[./-]\d{2,4}", line)) or bool(
            re.search(r"(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s+\d{4}", line, re.I)
        )
        has_event = any(x in lower for x in [
            "diagnosis", "started", "completed", "progression", "stable", "response",
            "scan", "chemotherapy", "surgery", "radiation", "admitted", "discharged", "follow"
        ])

        if has_date and has_event:
            events.append({
                "event_date": parse_possible_date(line),
                "event_type": infer_event_type(line),
                "title": line[:100],
                "details": line,
                "confidence": "low"
            })

    return events[:80]


def infer_event_type(line: str) -> str:
    lower = line.lower()
    if "progression" in lower:
        return "progression"
    if "stable" in lower:
        return "stable_disease"
    if "response" in lower:
        return "response"
    if "scan" in lower or "pet" in lower or "ct" in lower or "mri" in lower:
        return "scan"
    if "admitted" in lower or "admission" in lower:
        return "admission"
    if "discharge" in lower:
        return "discharge"
    if any(x in lower for x in ["chemo", "trastuzumab", "carboplatin", "cycle", "treatment"]):
        return "treatment"
    if "lab" in lower or "cbc" in lower or "lft" in lower or "kft" in lower:
        return "lab"
    return "other"


def fallback_scans(text: str) -> List[Dict[str, Any]]:
    scans = []
    lower = text.lower()

    if not any(x in lower for x in ["pet", "ct scan", "mri", "ultrasound", "x-ray", "xray", "fdg"]):
        return []

    blocks = re.split(r"(?=(?:PET[- ]?CT|CT Scan|CECT|MRI|Ultrasound|USG|X[- ]?Ray|Xray))", text, flags=re.I)

    for block in blocks:
        if not any(x in block.lower() for x in ["pet", "ct scan", "cect", "mri", "ultrasound", "usg", "x-ray", "xray"]):
            continue

        disease_status = "unclear"
        b = block.lower()

        if "progressive disease" in b or "disease progression" in b:
            disease_status = "progressive_disease"
        elif "stable disease" in b:
            disease_status = "stable_disease"
        elif "complete response" in b:
            disease_status = "complete_response"
        elif "partial response" in b:
            disease_status = "partial_response"
        elif "no scan evidence" in b or "absence of fdg avid disease" in b:
            disease_status = "no_active_disease"
        elif "mixed response" in b:
            disease_status = "mixed_response"

        scans.append({
            "scan_date": parse_possible_date(block) or regex_first(block, [r"\(([^)]+)\)"]),
            "scan_type": infer_scan_type(block),
            "body_region": None,
            "findings": compact_text(block, 1800),
            "comparison": regex_first(block, [r"As compared[^.]+\."]),
            "impression": disease_status.replace("_", " "),
            "disease_status": disease_status,
            "confidence": "low"
        })

    return scans[:40]


def infer_scan_type(text: str) -> Optional[str]:
    lower = text.lower()
    if "pet" in lower:
        return "PET-CT"
    if "cect" in lower:
        return "CECT"
    if "ct scan" in lower:
        return "CT"
    if "mri" in lower:
        return "MRI"
    if "ultrasound" in lower or "usg" in lower:
        return "Ultrasound"
    if "x-ray" in lower or "xray" in lower:
        return "X-ray"
    return None


def fallback_labs(text: str) -> List[Dict[str, Any]]:
    labs = []
    common_tests = [
        "hemoglobin", "haemoglobin", "hb", "wbc", "tlc", "platelet", "neutrophil",
        "lymphocyte", "creatinine", "urea", "bilirubin", "sgot", "sgpt", "alt", "ast",
        "alkaline phosphatase", "albumin", "sodium", "potassium", "calcium", "crp"
    ]

    for line in text.splitlines():
        lower = line.lower()
        if any(t in lower for t in common_tests):
            value_match = re.search(r"([-+]?\d+(?:\.\d+)?)", line)
            labs.append({
                "test_date": parse_possible_date(text),
                "panel_name": None,
                "test_name": compact_text(line, 120),
                "value": value_match.group(1) if value_match else None,
                "unit": None,
                "reference_range": None,
                "flag": "unknown",
                "clinical_note": compact_text(line, 500),
                "confidence": "low"
            })

    return labs[:100]


def fallback_medications(text: str) -> List[Dict[str, Any]]:
    meds = []

    for line in text.splitlines():
        clean = compact_text(line, 800)
        lower = clean.lower()

        if not any(x in lower for x in ["tab.", "cap.", "inj.", "syp.", "tablet", "capsule", "injection", "mouth wash", "ors"]):
            continue

        meds.append({
            "medicine_name": extract_med_name(clean),
            "dose": extract_dose(clean),
            "route": infer_route(clean),
            "frequency": extract_frequency(clean),
            "duration": extract_duration(clean),
            "timing": extract_timing(clean),
            "purpose": None,
            "start_date": parse_possible_date(clean),
            "end_date": None,
            "special_instructions": clean,
            "confidence": "low"
        })

    return meds[:120]


def extract_med_name(line: str) -> str:
    text = re.sub(r"^\s*[-•*]\s*", "", line).strip()
    return text[:140]


def extract_dose(line: str) -> Optional[str]:
    match = re.search(r"\b\d+(?:\.\d+)?\s*(?:mg|mcg|g|gm|ml|iu|units?|ampoules?|ampule|tabs?|caps?)\b", line, re.I)
    return match.group(0) if match else None


def infer_route(line: str) -> Optional[str]:
    lower = line.lower()
    if "subcutaneous" in lower or "s/c" in lower:
        return "subcutaneous"
    if " iv " in f" {lower} " or "intravenous" in lower:
        return "IV"
    if "mouth wash" in lower or "mouthwash" in lower:
        return "mouthwash"
    if "local application" in lower or "apply" in lower:
        return "local application"
    if lower.startswith("tab") or "tablet" in lower or lower.startswith("cap") or "capsule" in lower:
        return "oral"
    if lower.startswith("syp") or "syrup" in lower:
        return "oral"
    if lower.startswith("inj") or "injection" in lower:
        return "injection"
    return None


def extract_frequency(line: str) -> Optional[str]:
    lower = line.lower()
    patterns = [
        "once daily", "twice daily", "thrice daily", "three times daily",
        "once a day", "twice a day", "three times a day", "at bed time",
        "bedtime", "as needed", "as and when required", "after each episode",
        "every day", "weekly", "3 weekly", "q3 weekly", "q 3 weekly"
    ]
    found = [p for p in patterns if p in lower]
    return ", ".join(found) if found else None


def extract_duration(line: str) -> Optional[str]:
    patterns = [
        r"x\s*\d+\s*(?:days?|weeks?|months?)",
        r"for\s+\d+\s*(?:days?|weeks?|months?)",
        r"\d+\s*days?\s*on\s*/?\s*\d+\s*days?\s*off",
    ]
    for pat in patterns:
        m = re.search(pat, line, re.I)
        if m:
            return m.group(0)
    return None


def extract_timing(line: str) -> Optional[str]:
    lower = line.lower()
    timings = []
    for item in [
        "before meals", "after meals", "empty stomach", "after chemotherapy",
        "after 24 hours", "before injection", "before inj", "after first loose stool",
        "after each episode", "at bed time", "morning", "night", "evening"
    ]:
        if item in lower:
            timings.append(item)
    return ", ".join(timings) if timings else None


def fallback_protocols(text: str) -> List[Dict[str, Any]]:
    protocols = []
    lower = text.lower()

    protocol_triggers = [
        ("Loose stools / diarrhea", ["loose stool", "loose stools", "diarrhea", "diarrhoea"]),
        ("Nausea / vomiting", ["nausea", "vomiting"]),
        ("Oral ulcers / mouth care", ["oral ulcer", "mouth ulcer", "mouth wash", "mouthwash"]),
        ("Constipation", ["constipation"]),
        ("Pain", ["pain"]),
    ]

    for name, terms in protocol_triggers:
        if any(t in lower for t in terms):
            related_lines = []
            for line in text.splitlines():
                if any(t in line.lower() for t in terms):
                    related_lines.append(compact_text(line, 500))

            protocols.append({
                "protocol_name": name,
                "trigger_condition": ", ".join(terms),
                "steps": related_lines[:10],
                "escalation_rule": "Escalate to treating doctor/hospital if symptoms are severe, worsening, persistent, or listed as urgent in the report.",
                "confidence": "low"
            })

    return protocols


def fallback_urgent_rules(text: str) -> List[Dict[str, Any]]:
    rules = []

    urgent_section_patterns = [
        r"Please Seek Expert Medical Advice If(.*)",
        r"Report back to your Doctor(.*)",
        r"Emergency(.*)",
    ]

    section_text = ""
    for pat in urgent_section_patterns:
        m = re.search(pat, text, re.I | re.DOTALL)
        if m:
            section_text += "\n" + m.group(1)[:3000]

    if not section_text:
        section_text = text

    urgent_terms = [
        "high fever",
        "fever",
        "loose stools",
        "coffee coloured vomiting",
        "coffee colored vomiting",
        "black stools",
        "bleeding",
        "chest pain",
        "breathing difficulty",
        "loss of consciousness",
        "profuse sweating",
        "giddiness",
        "palpitation",
        "pain in abdomen",
        "reduced urine output",
        "severe weakness",
        "severe mouth ulcers",
        "rash",
        "swelling",
    ]

    for term in urgent_terms:
        if term in section_text.lower():
            rules.append({
                "risk_name": term.title(),
                "trigger_text": term,
                "severity": "urgent",
                "recommended_action": "Contact treating doctor/hospital or seek expert medical advice as instructed in the report.",
                "confidence": "low"
            })

    return rules


def fallback_contacts(text: str) -> List[Dict[str, Any]]:
    contacts = []
    phones = re.findall(r"(?:\+91[\s-]?)?[6-9]\d{9}", text)

    for phone in phones[:20]:
        contacts.append({
            "name": None,
            "role": "Phone number found in document",
            "phone": phone,
            "email": None,
            "notes": "Verify name/role against the original report before use.",
            "confidence": "low"
        })

    emails = re.findall(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    for email in emails[:10]:
        contacts.append({
            "name": None,
            "role": "Email found in document",
            "phone": None,
            "email": email,
            "notes": "Verify name/role against the original report before use.",
            "confidence": "low"
        })

    return contacts


def fallback_tasks(text: str) -> List[Dict[str, Any]]:
    tasks = []
    lower = text.lower()

    if "next cycle" in lower:
        tasks.append({
            "task_date": parse_possible_date(regex_first(text, [r"next cycle[^0-9]*(\d{1,2}[./-]\d{1,2}[./-]\d{2,4})"]) or ""),
            "task_type": "appointment",
            "title": "Next treatment cycle",
            "details": regex_first(text, [r"(Next Cycle[^\n]+)"]) or "Next cycle mentioned in document.",
            "confidence": "low"
        })

    if any(x in lower for x in ["cbc", "lft", "kft"]):
        tasks.append({
            "task_date": None,
            "task_type": "test",
            "title": "Prepare required blood reports",
            "details": "Document mentions blood reports/tests such as CBC/LFT/KFT. Verify required tests and timing with treating team.",
            "confidence": "low"
        })

    if "pet ct" in lower and "after" in lower:
        tasks.append({
            "task_date": None,
            "task_type": "scan",
            "title": "Planned scan follow-up",
            "details": regex_first(text, [r"(PET\s*CT[^\n]+)"]) or "PET/CT follow-up mentioned.",
            "confidence": "low"
        })

    return tasks


# ============================================================
# MERGE LAYER
# ============================================================

def save_document_record(file_name: str, file_hash_value: str, raw_text: str, classification: Dict[str, Any], extraction: Dict[str, Any]):
    execute("""
        INSERT OR REPLACE INTO documents
        (file_name, file_hash, document_type, confidence, patient_name, document_date, primary_purpose,
         raw_text, classification_json, extraction_json, uploaded_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        file_name,
        file_hash_value,
        classification.get("document_type"),
        classification.get("confidence"),
        classification.get("patient_name"),
        classification.get("document_date"),
        classification.get("primary_purpose"),
        raw_text,
        json.dumps(classification, ensure_ascii=False),
        json.dumps(extraction, ensure_ascii=False),
        now_iso()
    ))


def merge_case_profile(profile: Dict[str, Any], source: str):
    if not isinstance(profile, dict):
        return

    for key, value in profile.items():
        if value in [None, "", [], {}]:
            continue

        new_value = as_text(value)
        profile_key = str(key)

        existing = execute(
            "SELECT profile_value FROM case_profile WHERE profile_key = ?",
            (profile_key,),
            fetch=True
        )

        if existing:
            existing_value = existing[0][0]

            if existing_value and new_value and existing_value.strip().lower() != new_value.strip().lower():
                # Patient identifiers should not silently overwrite without conflict tracking.
                if profile_key in ["patient_name", "uhid_or_patient_id", "sex"]:
                    execute("""
                        INSERT INTO conflicts
                        (conflict_type, field_name, existing_value, new_value, source_document, status, created_at)
                        VALUES (?, ?, ?, ?, ?, ?, ?)
                    """, (
                        "case_profile_mismatch",
                        profile_key,
                        existing_value,
                        new_value,
                        source,
                        "open",
                        now_iso()
                    ))
                    continue

        execute("""
            INSERT INTO case_profile
            (profile_key, profile_value, source_document, confidence, updated_at)
            VALUES (?, ?, ?, ?, ?)
            ON CONFLICT(profile_key) DO UPDATE SET
                profile_value = excluded.profile_value,
                source_document = excluded.source_document,
                confidence = excluded.confidence,
                updated_at = excluded.updated_at
        """, (
            profile_key,
            new_value,
            source,
            "medium",
            now_iso()
        ))


def merge_diagnosis(records: List[Dict[str, Any]], source: str):
    for r in records:
        if not isinstance(r, dict):
            continue

        if not any(r.get(k) for k in ["primary_diagnosis", "cancer_type", "site", "histology", "stage", "biomarkers", "details"]):
            continue

        dedupe = dedupe_hash(
            "diagnosis",
            r.get("diagnosis_date"),
            r.get("primary_diagnosis"),
            r.get("site"),
            r.get("stage"),
            source
        )

        insert_ignore("""
            INSERT INTO diagnosis_records
            (diagnosis_date, primary_diagnosis, cancer_type, site, histology, stage, grade,
             metastatic_status, biomarkers_json, details, source_document, confidence, dedupe_key, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            parse_possible_date(r.get("diagnosis_date")) or r.get("diagnosis_date"),
            r.get("primary_diagnosis"),
            r.get("cancer_type"),
            r.get("site"),
            r.get("histology"),
            r.get("stage"),
            r.get("grade"),
            r.get("metastatic_status"),
            json.dumps(as_list(r.get("biomarkers")), ensure_ascii=False),
            r.get("details"),
            source,
            r.get("confidence") or "medium",
            dedupe,
            now_iso()
        ))


def merge_treatment_plans(plans: List[Dict[str, Any]], source: str):
    for p in plans:
        if not isinstance(p, dict):
            continue

        if not any(p.get(k) for k in ["regimen_name", "drugs", "cycle_info", "next_cycle_date", "details", "scan_plan"]):
            continue

        dedupe = dedupe_hash(
            "treatment",
            p.get("plan_date"),
            p.get("regimen_name"),
            p.get("cycle_info"),
            p.get("next_cycle_date"),
            source
        )

        insert_ignore("""
            INSERT INTO treatment_plans
            (plan_date, treatment_line, regimen_name, drugs_json, intent, cycle_info,
             next_cycle_date, required_tests_json, scan_plan, details, source_document,
             confidence, dedupe_key, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            parse_possible_date(p.get("plan_date")) or p.get("plan_date"),
            p.get("treatment_line"),
            p.get("regimen_name"),
            json.dumps(as_list(p.get("drugs")), ensure_ascii=False),
            p.get("intent"),
            p.get("cycle_info"),
            parse_possible_date(p.get("next_cycle_date")) or p.get("next_cycle_date"),
            json.dumps(as_list(p.get("required_tests")), ensure_ascii=False),
            p.get("scan_plan"),
            p.get("details"),
            source,
            p.get("confidence") or "medium",
            dedupe,
            now_iso()
        ))


def merge_timeline(events: List[Dict[str, Any]], source: str):
    for e in events:
        if not isinstance(e, dict):
            continue

        if not any(e.get(k) for k in ["title", "details", "event_date"]):
            continue

        dedupe = dedupe_hash(
            "timeline",
            e.get("event_date"),
            e.get("event_type"),
            e.get("title"),
            e.get("details")
        )

        insert_ignore("""
            INSERT INTO timeline_events
            (event_date, event_type, title, details, source_document, confidence, dedupe_key, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            parse_possible_date(e.get("event_date")) or e.get("event_date"),
            e.get("event_type"),
            e.get("title"),
            e.get("details"),
            source,
            e.get("confidence") or "medium",
            dedupe,
            now_iso()
        ))


def merge_scans(scans: List[Dict[str, Any]], source: str):
    for s in scans:
        if not isinstance(s, dict):
            continue

        if not any(s.get(k) for k in ["scan_date", "scan_type", "findings", "impression"]):
            continue

        dedupe = dedupe_hash(
            "scan",
            s.get("scan_date"),
            s.get("scan_type"),
            s.get("body_region"),
            s.get("impression"),
            compact_text(s.get("findings"), 400)
        )

        insert_ignore("""
            INSERT INTO scan_events
            (scan_date, scan_type, body_region, findings, comparison, impression,
             disease_status, source_document, confidence, dedupe_key, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            parse_possible_date(s.get("scan_date")) or s.get("scan_date"),
            s.get("scan_type"),
            s.get("body_region"),
            s.get("findings"),
            s.get("comparison"),
            s.get("impression"),
            s.get("disease_status"),
            source,
            s.get("confidence") or "medium",
            dedupe,
            now_iso()
        ))


def merge_labs(labs: List[Dict[str, Any]], source: str):
    for l in labs:
        if not isinstance(l, dict):
            continue

        if not any(l.get(k) for k in ["test_name", "value", "panel_name"]):
            continue

        dedupe = dedupe_hash(
            "lab",
            l.get("test_date"),
            l.get("panel_name"),
            l.get("test_name"),
            l.get("value"),
            source
        )

        insert_ignore("""
            INSERT INTO lab_results
            (test_date, panel_name, test_name, value, unit, reference_range, flag,
             clinical_note, source_document, confidence, dedupe_key, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            parse_possible_date(l.get("test_date")) or l.get("test_date"),
            l.get("panel_name"),
            l.get("test_name"),
            l.get("value"),
            l.get("unit"),
            l.get("reference_range"),
            l.get("flag"),
            l.get("clinical_note"),
            source,
            l.get("confidence") or "medium",
            dedupe,
            now_iso()
        ))


def merge_medications(meds: List[Dict[str, Any]], source: str):
    for m in meds:
        if not isinstance(m, dict):
            continue

        if not any(m.get(k) for k in ["medicine_name", "dose", "frequency", "special_instructions"]):
            continue

        dedupe = dedupe_hash(
            "med",
            m.get("medicine_name"),
            m.get("dose"),
            m.get("frequency"),
            m.get("duration"),
            m.get("start_date"),
            source
        )

        insert_ignore("""
            INSERT INTO medications
            (medicine_name, dose, route, frequency, duration, timing, purpose,
             start_date, end_date, special_instructions, source_document, confidence,
             dedupe_key, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            m.get("medicine_name"),
            m.get("dose"),
            m.get("route"),
            m.get("frequency"),
            m.get("duration"),
            m.get("timing"),
            m.get("purpose"),
            parse_possible_date(m.get("start_date")) or m.get("start_date"),
            parse_possible_date(m.get("end_date")) or m.get("end_date"),
            m.get("special_instructions"),
            source,
            m.get("confidence") or "medium",
            dedupe,
            now_iso()
        ))


def merge_protocols(protocols: List[Dict[str, Any]], source: str):
    for p in protocols:
        if not isinstance(p, dict):
            continue

        if not any(p.get(k) for k in ["protocol_name", "trigger_condition", "steps", "escalation_rule"]):
            continue

        dedupe = dedupe_hash(
            "protocol",
            p.get("protocol_name"),
            p.get("trigger_condition"),
            source
        )

        insert_ignore("""
            INSERT INTO side_effect_protocols
            (protocol_name, trigger_condition, steps_json, escalation_rule,
             source_document, confidence, dedupe_key, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            p.get("protocol_name"),
            p.get("trigger_condition"),
            json.dumps(as_list(p.get("steps")), ensure_ascii=False),
            p.get("escalation_rule"),
            source,
            p.get("confidence") or "medium",
            dedupe,
            now_iso()
        ))


def merge_urgent_rules(rules: List[Dict[str, Any]], source: str):
    for r in rules:
        if not isinstance(r, dict):
            continue

        if not any(r.get(k) for k in ["risk_name", "trigger_text", "recommended_action"]):
            continue

        dedupe = dedupe_hash(
            "urgent",
            r.get("risk_name"),
            r.get("trigger_text"),
            r.get("severity"),
            source
        )

        insert_ignore("""
            INSERT INTO urgent_rules
            (risk_name, trigger_text, severity, recommended_action,
             source_document, confidence, dedupe_key, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            r.get("risk_name"),
            r.get("trigger_text"),
            r.get("severity"),
            r.get("recommended_action"),
            source,
            r.get("confidence") or "medium",
            dedupe,
            now_iso()
        ))


def merge_contacts(contacts: List[Dict[str, Any]], source: str):
    for c in contacts:
        if not isinstance(c, dict):
            continue

        if not any(c.get(k) for k in ["name", "phone", "email"]):
            continue

        dedupe = dedupe_hash(
            "contact",
            c.get("name"),
            c.get("role"),
            c.get("phone"),
            c.get("email")
        )

        insert_ignore("""
            INSERT INTO doctor_contacts
            (name, role, phone, email, notes, source_document, confidence, dedupe_key, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            c.get("name"),
            c.get("role"),
            c.get("phone"),
            c.get("email"),
            c.get("notes"),
            source,
            c.get("confidence") or "medium",
            dedupe,
            now_iso()
        ))


def merge_tasks(tasks: List[Dict[str, Any]], source: str):
    for t in tasks:
        if not isinstance(t, dict):
            continue

        if not any(t.get(k) for k in ["title", "details", "task_date"]):
            continue

        dedupe = dedupe_hash(
            "task",
            t.get("task_date"),
            t.get("task_type"),
            t.get("title"),
            source
        )

        insert_ignore("""
            INSERT INTO care_tasks
            (task_date, task_type, title, details, status, source_document, confidence, dedupe_key, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            parse_possible_date(t.get("task_date")) or t.get("task_date"),
            t.get("task_type"),
            t.get("title"),
            t.get("details"),
            "pending",
            source,
            t.get("confidence") or "medium",
            dedupe,
            now_iso()
        ))


def merge_extraction_into_case(extraction: Dict[str, Any], source: str):
    classification = normalize_classification(extraction.get("document_classification", {}))
    should = classification.get("should_update", {})

    # Always save profile if available, but conflict-protect identifiers.
    if should.get("case_profile"):
        merge_case_profile(extraction.get("case_profile", {}), source)

    if should.get("diagnosis_records"):
        merge_diagnosis(extraction.get("diagnosis_records", []), source)

    if should.get("treatment_plans"):
        merge_treatment_plans(extraction.get("treatment_plans", []), source)

    if should.get("timeline_events"):
        merge_timeline(extraction.get("timeline_events", []), source)

    if should.get("scan_events"):
        merge_scans(extraction.get("scan_events", []), source)

    if should.get("lab_results"):
        merge_labs(extraction.get("lab_results", []), source)

    if should.get("medications"):
        merge_medications(extraction.get("medications", []), source)

    if should.get("side_effect_protocols"):
        merge_protocols(extraction.get("side_effect_protocols", []), source)

    if should.get("urgent_rules"):
        merge_urgent_rules(extraction.get("urgent_rules", []), source)

    if should.get("doctor_contacts"):
        merge_contacts(extraction.get("doctor_contacts", []), source)

    if should.get("care_tasks"):
        merge_tasks(extraction.get("care_tasks", []), source)


# ============================================================
# DATAFRAMES
# ============================================================

def df_query(query: str, columns: List[str]) -> pd.DataFrame:
    rows = execute(query, fetch=True)
    return pd.DataFrame(rows, columns=columns)


def get_documents_df():
    return df_query("""
        SELECT file_name, document_type, confidence, patient_name, document_date, primary_purpose, uploaded_at
        FROM documents
        ORDER BY uploaded_at DESC
    """, ["File", "Type", "Confidence", "Patient", "Date", "Purpose", "Uploaded At"])


def get_profile_df():
    return df_query("""
        SELECT profile_key, profile_value, source_document, confidence, updated_at
        FROM case_profile
        ORDER BY profile_key
    """, ["Key", "Value", "Source", "Confidence", "Updated At"])


def get_conflicts_df():
    return df_query("""
        SELECT conflict_type, field_name, existing_value, new_value, source_document, status, created_at
        FROM conflicts
        ORDER BY created_at DESC
    """, ["Type", "Field", "Existing", "New", "Source", "Status", "Created At"])


def get_diagnosis_df():
    rows = execute("""
        SELECT diagnosis_date, primary_diagnosis, cancer_type, site, histology, stage, grade,
               metastatic_status, biomarkers_json, details, source_document, confidence
        FROM diagnosis_records
        ORDER BY
            CASE WHEN diagnosis_date IS NULL THEN 1 ELSE 0 END,
            diagnosis_date ASC,
            created_at ASC
    """, fetch=True)

    parsed = []
    for r in rows:
        try:
            markers = ", ".join(json.loads(r[8] or "[]"))
        except Exception:
            markers = ""

        parsed.append(list(r[:8]) + [markers] + list(r[9:]))

    return pd.DataFrame(parsed, columns=[
        "Date", "Diagnosis", "Cancer Type", "Site", "Histology", "Stage", "Grade",
        "Metastatic Status", "Biomarkers", "Details", "Source", "Confidence"
    ])


def get_treatment_df():
    rows = execute("""
        SELECT plan_date, treatment_line, regimen_name, drugs_json, intent, cycle_info,
               next_cycle_date, required_tests_json, scan_plan, details, source_document, confidence
        FROM treatment_plans
        ORDER BY
            CASE WHEN plan_date IS NULL THEN 1 ELSE 0 END,
            plan_date ASC,
            created_at ASC
    """, fetch=True)

    parsed = []
    for r in rows:
        try:
            drugs = ", ".join(json.loads(r[3] or "[]"))
        except Exception:
            drugs = ""

        try:
            tests = ", ".join(json.loads(r[7] or "[]"))
        except Exception:
            tests = ""

        parsed.append([
            r[0], r[1], r[2], drugs, r[4], r[5], r[6], tests, r[8], r[9], r[10], r[11]
        ])

    return pd.DataFrame(parsed, columns=[
        "Plan Date", "Line", "Regimen", "Drugs", "Intent", "Cycle Info", "Next Cycle",
        "Required Tests", "Scan Plan", "Details", "Source", "Confidence"
    ])


def get_timeline_df():
    df = df_query("""
        SELECT event_date, event_type, title, details, source_document, confidence
        FROM timeline_events
        ORDER BY created_at ASC
    """, ["Date", "Type", "Title", "Details", "Source", "Confidence"])

    if df.empty:
        return df

    return make_timeline_display_df(df)

def get_scans_df():
    return df_query("""
        SELECT scan_date, scan_type, body_region, disease_status, impression, findings, comparison, source_document, confidence
        FROM scan_events
        ORDER BY
            CASE WHEN scan_date IS NULL THEN 1 ELSE 0 END,
            scan_date ASC,
            created_at ASC
    """, ["Date", "Scan Type", "Region", "Disease Status", "Impression", "Findings", "Comparison", "Source", "Confidence"])


def get_labs_df():
    return df_query("""
        SELECT test_date, panel_name, test_name, value, unit, reference_range, flag, clinical_note, source_document, confidence
        FROM lab_results
        ORDER BY
            CASE WHEN test_date IS NULL THEN 1 ELSE 0 END,
            test_date ASC,
            panel_name ASC,
            test_name ASC
    """, ["Date", "Panel", "Test", "Value", "Unit", "Reference Range", "Flag", "Note", "Source", "Confidence"])


def get_meds_df():
    return df_query("""
        SELECT medicine_name, dose, route, frequency, duration, timing, purpose,
               start_date, end_date, special_instructions, source_document, confidence
        FROM medications
        ORDER BY created_at ASC
    """, ["Medicine", "Dose", "Route", "Frequency", "Duration", "Timing", "Purpose", "Start", "End", "Instructions", "Source", "Confidence"])


def get_protocols_df():
    rows = execute("""
        SELECT protocol_name, trigger_condition, steps_json, escalation_rule, source_document, confidence
        FROM side_effect_protocols
        ORDER BY created_at ASC
    """, fetch=True)

    parsed = []
    for r in rows:
        try:
            steps = "\n".join([f"- {x}" for x in json.loads(r[2] or "[]")])
        except Exception:
            steps = r[2] or ""
        parsed.append([r[0], r[1], steps, r[3], r[4], r[5]])

    return pd.DataFrame(parsed, columns=["Protocol", "Trigger", "Steps", "Escalation", "Source", "Confidence"])


def get_urgent_df():
    return df_query("""
        SELECT risk_name, trigger_text, severity, recommended_action, source_document, confidence
        FROM urgent_rules
        ORDER BY
            CASE severity
                WHEN 'emergency' THEN 1
                WHEN 'urgent' THEN 2
                ELSE 3
            END,
            created_at ASC
    """, ["Risk", "Trigger", "Severity", "Action", "Source", "Confidence"])


def get_contacts_df():
    return df_query("""
        SELECT name, role, phone, email, notes, source_document, confidence
        FROM doctor_contacts
        ORDER BY created_at ASC
    """, ["Name", "Role", "Phone", "Email", "Notes", "Source", "Confidence"])


def get_tasks_df():
    return df_query("""
        SELECT id, task_date, task_type, title, details, status, source_document, confidence
        FROM care_tasks
        ORDER BY
            CASE WHEN task_date IS NULL THEN 1 ELSE 0 END,
            task_date ASC,
            created_at ASC
    """, ["ID", "Date", "Type", "Title", "Details", "Status", "Source", "Confidence"])


def get_symptoms_df():
    rows = execute("""
        SELECT symptom_date, symptom_name, severity, temperature, notes, urgent_flags_json, created_at
        FROM symptoms
        ORDER BY symptom_date DESC, created_at DESC
    """, fetch=True)

    parsed = []
    for r in rows:
        try:
            flags = "; ".join(json.loads(r[5] or "[]"))
        except Exception:
            flags = ""
        parsed.append([r[0], r[1], r[2], r[3], r[4], flags, r[6]])

    return pd.DataFrame(parsed, columns=["Date", "Symptom", "Severity", "Temperature", "Notes", "Urgent Flags", "Created At"])


def get_memory_df():
    return df_query("""
        SELECT memory_type, title, details, created_at
        FROM caregiver_memory
        ORDER BY created_at DESC
    """, ["Type", "Title", "Details", "Created At"])


def profile_value(key: str) -> str:
    rows = execute("SELECT profile_value FROM case_profile WHERE profile_key = ?", (key,), fetch=True)
    return rows[0][0] if rows else ""


# ============================================================
# SYMPTOM RISK ENGINE
# ============================================================

GENERIC_URGENT_RULES = [
    {
        "keywords": ["fever", "temperature", "high fever"],
        "severity_min": 1,
        "message": "Fever can be urgent in cancer patients, especially after chemotherapy. Contact the treating doctor/hospital."
    },
    {
        "keywords": ["breathing difficulty", "breathlessness", "shortness of breath"],
        "severity_min": 1,
        "message": "Breathing difficulty is urgent. Seek medical help."
    },
    {
        "keywords": ["chest pain"],
        "severity_min": 1,
        "message": "Chest pain should be escalated urgently."
    },
    {
        "keywords": ["loss of consciousness", "unconscious", "fainted"],
        "severity_min": 1,
        "message": "Loss of consciousness is an emergency sign."
    },
    {
        "keywords": ["black stool", "black stools", "tar", "coffee coloured vomiting", "coffee colored vomiting"],
        "severity_min": 1,
        "message": "Black stools or coffee-colored vomiting may suggest bleeding. Seek urgent medical advice."
    },
    {
        "keywords": ["bleeding", "blood"],
        "severity_min": 1,
        "message": "Bleeding from any site should be escalated urgently."
    },
    {
        "keywords": ["reduced urine", "low urine", "not passing urine"],
        "severity_min": 1,
        "message": "Reduced urine output can be a warning sign. Contact the doctor."
    },
    {
        "keywords": ["severe weakness", "very weak", "unable to walk"],
        "severity_min": 1,
        "message": "Severe weakness should be escalated."
    },
    {
        "keywords": ["rash", "swelling", "face swelling", "body swelling"],
        "severity_min": 1,
        "message": "Rash or swelling may indicate a reaction. Contact the treating team."
    },
    {
        "keywords": ["loose stool", "loose stools", "diarrhea", "diarrhoea"],
        "severity_min": 5,
        "message": "Frequent loose stools can cause dehydration. Follow the prescribed protocol and escalate if persistent."
    },
    {
        "keywords": ["vomiting", "nausea"],
        "severity_min": 7,
        "message": "Persistent vomiting can cause dehydration. Contact the treating team if not controlled."
    },
]


def extract_temp(text: str) -> Optional[float]:
    if not text:
        return None

    lower = text.lower()
    matches = re.findall(r"(\d{2,3}(?:\.\d+)?)\s*(c|f|°c|°f)?", lower)

    for val, unit in matches:
        try:
            num = float(val)
        except Exception:
            continue

        if unit in ["f", "°f"] and num >= 100.4:
            return num

        if unit in ["c", "°c"] and num >= 38.0:
            return num

        if not unit:
            if 38 <= num <= 45:
                return num
            if num >= 100.4:
                return num

    return None


def evaluate_symptom(symptom: str, severity: int, temperature: str, notes: str) -> List[str]:
    combined = f"{symptom} {temperature} {notes}".lower()
    flags = []

    if extract_temp(combined) is not None:
        flags.append("Fever threshold detected. Contact the treating doctor/hospital, especially if chemotherapy was recent.")

    for rule in GENERIC_URGENT_RULES:
        if severity >= rule["severity_min"] and any(k in combined for k in rule["keywords"]):
            flags.append(rule["message"])

    custom_rules = get_urgent_df()
    if not custom_rules.empty:
        for _, row in custom_rules.iterrows():
            trigger = str(row["Trigger"] or "").lower()
            risk = str(row["Risk"] or "").lower()
            action = str(row["Action"] or "")

            if trigger and trigger in combined:
                flags.append(action)

            elif risk and risk in combined:
                flags.append(action)

    return sorted(list(set([f for f in flags if f])))


def save_symptom(symptom_date: str, symptom_name: str, severity: int, temperature: str, notes: str, flags: List[str]):
    execute("""
        INSERT INTO symptoms
        (symptom_date, symptom_name, severity, temperature, notes, urgent_flags_json, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (
        symptom_date,
        symptom_name,
        severity,
        temperature,
        notes,
        json.dumps(flags, ensure_ascii=False),
        now_iso()
    ))


# ============================================================
# CASE BRIEF
# ============================================================

def build_case_context() -> str:
    sections = []

    frames = {
        "CASE PROFILE": get_profile_df(),
        "DIAGNOSIS": get_diagnosis_df(),
        "TREATMENT PLANS": get_treatment_df(),
        "TIMELINE": get_timeline_df(),
        "SCANS": get_scans_df(),
        "LABS": get_labs_df(),
        "MEDICATIONS": get_meds_df(),
        "PROTOCOLS": get_protocols_df(),
        "URGENT RULES": get_urgent_df(),
        "TASKS": get_tasks_df(),
        "SYMPTOMS": get_symptoms_df(),
        "CAREGIVER MEMORY": get_memory_df(),
        "CONFLICTS": get_conflicts_df(),
    }

    for name, df in frames.items():
        if df is not None and not df.empty:
            sections.append(f"\n\n{name}\n{df.head(150).to_string(index=False)}")

    return "\n".join(sections)


def generate_case_brief() -> str:
    context = build_case_context()

    if not context.strip():
        return "No case data available yet. Upload reports first."

    if not client:
        return "AI brief unavailable because OPENAI_API_KEY is not configured."

    prompt = f"""
Create a caregiver-ready case brief from the patient case context.

Format:

# Caregiver Case Brief

## 1. Current Known Situation
Simple, factual, from records.

## 2. Disease / Treatment Timeline
Key events in order.

## 3. Latest Scan / Lab / Treatment Updates
Separate each if present.

## 4. Medicines and Supportive Care
Do not alter doses. Say verify with original prescription.

## 5. What To Monitor At Home
Practical monitoring checklist.

## 6. Urgent Red Flags
Clear escalation list.

## 7. Upcoming Tasks
Tests, reports, appointments, scans, medicines.

## 8. Questions For Doctor
Sharp and practical.

## 9. Missing / Unclear / Conflicting Items
Mention conflicts or uncertain data.

Rules:
- Do not diagnose.
- Do not prescribe.
- Do not recommend treatment changes.
- Do not invent anything.
- Clearly say when something is unclear.
- Final guidance must come from treating doctor.

Context:
{context[:65000]}
"""

    response = client.chat.completions.create(
        model=OPENAI_MODEL,
        temperature=0.15,
        messages=[
            {"role": "system", "content": "You summarize structured cancer-care case data for family caregivers."},
            {"role": "user", "content": prompt}
        ],
    )

    return response.choices[0].message.content


def export_case_json() -> str:
    payload = {
        "documents": get_documents_df().to_dict(orient="records"),
        "case_profile": get_profile_df().to_dict(orient="records"),
        "conflicts": get_conflicts_df().to_dict(orient="records"),
        "diagnosis": get_diagnosis_df().to_dict(orient="records"),
        "treatment_plans": get_treatment_df().to_dict(orient="records"),
        "timeline": get_timeline_df().to_dict(orient="records"),
        "scans": get_scans_df().to_dict(orient="records"),
        "labs": get_labs_df().to_dict(orient="records"),
        "medications": get_meds_df().to_dict(orient="records"),
        "protocols": get_protocols_df().to_dict(orient="records"),
        "urgent_rules": get_urgent_df().to_dict(orient="records"),
        "contacts": get_contacts_df().to_dict(orient="records"),
        "tasks": get_tasks_df().to_dict(orient="records"),
        "symptoms": get_symptoms_df().to_dict(orient="records"),
        "caregiver_memory": get_memory_df().to_dict(orient="records"),
        "exported_at": now_iso(),
    }
    return json.dumps(payload, ensure_ascii=False, indent=2)


# ============================================================
# STREAMLIT UI HELPERS
# ============================================================

def render_df(df: pd.DataFrame, empty: str):
    if df.empty:
        st.info(empty)
    else:
        st.dataframe(df, use_container_width=True, hide_index=True)


def metric(label: str, value: Any):
    st.metric(label, value if value not in [None, "", [], {}] else "—")


def warning_banner():
    st.warning(
        "This tool organizes medical information for caregivers. It does not diagnose, prescribe, or replace the treating doctor. "
        "For emergency symptoms, contact the treating hospital/doctor or emergency services."
    )


def summarize_extraction_counts(extraction: Dict[str, Any]) -> Dict[str, int]:
    keys = [
        "diagnosis_records",
        "treatment_plans",
        "timeline_events",
        "scan_events",
        "lab_results",
        "medications",
        "side_effect_protocols",
        "urgent_rules",
        "doctor_contacts",
        "care_tasks",
    ]
    return {k: len(as_list(extraction.get(k))) for k in keys}


def render_processing_result(file_name: str, classification: Dict[str, Any], extraction: Dict[str, Any]):
    st.success(f"Processed: {file_name}")

    c1, c2, c3 = st.columns(3)
    with c1:
        metric("Document Type", classification.get("document_type"))
    with c2:
        metric("Confidence", classification.get("confidence"))
    with c3:
        metric("Document Date", classification.get("document_date"))

    if classification.get("primary_purpose"):
        st.info(f"**Purpose:** {classification.get('primary_purpose')}")

    counts = summarize_extraction_counts(extraction)

    st.markdown("### Updated Case Areas")
    updates = classification.get("should_update", {})

    rows = []
    for area, should in updates.items():
        matching_key = area
        count = counts.get(matching_key, None)
        rows.append({
            "Area": area,
            "Should Update": "Yes" if should else "No",
            "Extracted Items": count if count is not None else ""
        })

    st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)

    explanation = extraction.get("caregiver_explanation", {})
    if isinstance(explanation, dict):
        with st.expander("Caregiver Explanation From This Document", expanded=True):
            if explanation.get("simple_summary"):
                st.write(explanation.get("simple_summary"))
            if explanation.get("what_changed"):
                st.markdown(f"**What changed:** {explanation.get('what_changed')}")

            next_items = as_list(explanation.get("what_to_do_next"))
            if next_items:
                st.markdown("**What to do next:**")
                for item in next_items:
                    st.markdown(f"- {item}")

            questions = as_list(explanation.get("questions_for_doctor"))
            if questions:
                st.markdown("**Questions for doctor:**")
                for q in questions:
                    st.markdown(f"- {q}")

            missing = as_list(explanation.get("missing_or_unclear_items"))
            if missing:
                st.markdown("**Missing / unclear:**")
                for m in missing:
                    st.markdown(f"- {m}")


# ============================================================
# TABS
# ============================================================

def tab_upload():
    st.header("Upload Reports")

    st.write(
        "Upload reports one by one or together. The app will classify each document first, then update only the relevant parts of the patient case."
    )

    uploaded_files = st.file_uploader(
        "Upload PDF, DOCX, or TXT",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True
    )

    force = st.checkbox("Force reprocess existing files", value=False)

    if uploaded_files:
        for uploaded_file in uploaded_files:
            file_bytes = uploaded_file.read()
            h = file_hash(file_bytes)

            existing = execute("SELECT id FROM documents WHERE file_hash = ?", (h,), fetch=True)
            if existing and not force:
                st.info(f"Already processed: {uploaded_file.name}")
                continue

            with st.spinner(f"Processing {uploaded_file.name}..."):
                try:
                    raw_text = extract_text_from_file(uploaded_file.name, file_bytes)

                    if len(raw_text.strip()) < 80:
                        st.error(
                            f"Very little text extracted from {uploaded_file.name}. "
                            "This may be a scanned image PDF. Use a clearer PDF or add OCR support."
                        )
                        continue

                    classification = classify_document_ai(raw_text, uploaded_file.name)
                    extraction = ai_extract_document(raw_text, uploaded_file.name, classification)

                    # Ensure classification used for merge is the top-level normalized one
                    classification = normalize_classification(extraction.get("document_classification") or classification)
                    extraction["document_classification"] = classification

                    save_document_record(uploaded_file.name, h, raw_text, classification, extraction)
                    merge_extraction_into_case(extraction, uploaded_file.name)

                    render_processing_result(uploaded_file.name, classification, extraction)

                except Exception as e:
                    st.error(f"Failed to process {uploaded_file.name}: {e}")

    st.subheader("Uploaded Documents")
    render_df(get_documents_df(), "No documents uploaded yet.")


def tab_dashboard():
    st.header("Unified Patient Case")

    patient_name = profile_value("patient_name")
    patient_id = profile_value("uhid_or_patient_id")
    consultant = profile_value("consultant")
    allergy = profile_value("allergies")

    diagnosis_df = get_diagnosis_df()
    treatment_df = get_treatment_df()
    scan_df = get_scans_df()
    lab_df = get_labs_df()
    meds_df = get_meds_df()
    urgent_df = get_urgent_df()
    tasks_df = get_tasks_df()
    conflicts_df = get_conflicts_df()

    c1, c2, c3, c4 = st.columns(4)
    with c1:
        metric("Patient", patient_name)
    with c2:
        metric("Patient ID", patient_id)
    with c3:
        metric("Consultant", consultant)
    with c4:
        metric("Allergies", allergy)

    c5, c6, c7, c8 = st.columns(4)
    with c5:
        metric("Diagnosis Records", len(diagnosis_df))
    with c6:
        metric("Scans", len(scan_df))
    with c7:
        metric("Lab Results", len(lab_df))
    with c8:
        metric("Medicines", len(meds_df))

    if not conflicts_df.empty:
        st.error("Open conflicts found. Review patient identifier/profile mismatches.")
        render_df(conflicts_df, "No conflicts.")

    st.subheader("Current / Latest Treatment Plans")
    render_df(treatment_df.tail(5), "No treatment plan extracted yet.")

    st.subheader("Pending Tasks")
    if not tasks_df.empty:
        pending = tasks_df[tasks_df["Status"] == "pending"]
        render_df(pending, "No pending tasks.")
    else:
        st.info("No tasks extracted yet.")

    st.subheader("Urgent Red Flags")
    render_df(urgent_df, "No urgent red flags extracted yet.")


def tab_profile():
    st.header("Patient Profile")

    render_df(get_profile_df(), "No patient profile yet.")

    st.subheader("Diagnosis Records")
    render_df(get_diagnosis_df(), "No diagnosis records yet.")

    st.subheader("Doctor / Hospital Contacts")
    render_df(get_contacts_df(), "No contacts extracted yet.")

    st.subheader("Conflicts")
    render_df(get_conflicts_df(), "No conflicts found.")


def tab_timeline():
    st.header("Patient Timeline")

    df = get_timeline_df()

    if df.empty:
        st.info("No timeline events yet.")
        return

    st.info(
        "Timeline is grouped into Current / Latest, Chronological, and Date Unclear. "
        "Older history repeated inside newer summaries is treated as historical context, not current status."
    )

    buckets = ["Current / Latest", "Chronological", "Date Unclear"]
    selected_buckets = st.multiselect(
        "Filter by timeline bucket",
        buckets,
        default=buckets
    )

    event_types = sorted([x for x in df["Type"].dropna().unique().tolist() if x])
    selected_types = st.multiselect("Filter by event type", event_types)

    filtered = df.copy()

    if selected_buckets:
        filtered = filtered[filtered["Bucket"].isin(selected_buckets)]

    if selected_types:
        filtered = filtered[filtered["Type"].isin(selected_types)]

    display_cols = [
        "Bucket",
        "Display Date",
        "Type",
        "Title",
        "Details",
        "Source",
        "Confidence",
    ]

    st.subheader("Timeline Table")
    render_df(filtered[display_cols], "No matching timeline events.")

    st.subheader("Timeline Cards")

    for bucket in buckets:
        section = filtered[filtered["Bucket"] == bucket]

        if section.empty:
            continue

        st.markdown(f"## {bucket}")

        for _, row in section.iterrows():
            st.markdown(f"### {row['Display Date']} — {row['Title'] or 'Event'}")
            st.markdown(
                f"**Type:** {row['Type']}  \n"
                f"**Source:** {row['Source']}  \n"
                f"**Confidence:** {row['Confidence']}"
            )

            if row["Details"]:
                st.write(row["Details"])

            st.divider()

def tab_scans():
    st.header("Scan Timeline")

    df = get_scans_df()
    if df.empty:
        st.info("No scan events yet.")
        return

    statuses = sorted([x for x in df["Disease Status"].dropna().unique().tolist() if x])
    selected = st.multiselect("Filter by disease status", statuses)

    filtered = df.copy()
    if selected:
        filtered = filtered[filtered["Disease Status"].isin(selected)]

    render_df(filtered, "No matching scans.")

    st.subheader("Scan Cards")
    for _, row in filtered.iterrows():
        st.markdown(f"### {row['Date'] or 'Date unclear'} — {row['Scan Type'] or 'Scan'}")
        st.markdown(f"**Region:** {row['Region'] or 'Not specified'}")
        st.markdown(f"**Disease Status:** {row['Disease Status'] or 'Unclear'}")
        if row["Impression"]:
            st.markdown(f"**Impression:** {row['Impression']}")
        if row["Comparison"]:
            st.markdown(f"**Comparison:** {row['Comparison']}")
        st.write(row["Findings"])
        st.caption(f"Source: {row['Source']} | Confidence: {row['Confidence']}")
        st.divider()


def tab_labs():
    st.header("Lab Results")

    df = get_labs_df()
    if df.empty:
        st.info("No lab results yet.")
        return

    panels = sorted([x for x in df["Panel"].dropna().unique().tolist() if x])
    flags = sorted([x for x in df["Flag"].dropna().unique().tolist() if x])

    c1, c2 = st.columns(2)
    with c1:
        selected_panels = st.multiselect("Filter by panel", panels)
    with c2:
        selected_flags = st.multiselect("Filter by flag", flags)

    filtered = df.copy()
    if selected_panels:
        filtered = filtered[filtered["Panel"].isin(selected_panels)]
    if selected_flags:
        filtered = filtered[filtered["Flag"].isin(selected_flags)]

    render_df(filtered, "No matching lab results.")


def tab_medicines():
    st.header("Medicines")

    st.error("Verify all medicine names, doses, timings, and duration against the original prescription or treating doctor before use.")

    df = get_meds_df()
    if df.empty:
        st.info("No medicines extracted yet.")
        return

    purposes = sorted([x for x in df["Purpose"].dropna().unique().tolist() if x])
    selected = st.multiselect("Filter by purpose", purposes)

    filtered = df.copy()
    if selected:
        filtered = filtered[filtered["Purpose"].isin(selected)]

    render_df(filtered, "No matching medicines.")

    st.subheader("Medicine Cards")
    for _, row in filtered.iterrows():
        st.markdown(f"### {row['Medicine']}")
        c1, c2, c3, c4 = st.columns(4)
        with c1:
            st.markdown(f"**Dose:** {row['Dose'] or 'Verify'}")
        with c2:
            st.markdown(f"**Frequency:** {row['Frequency'] or 'Verify'}")
        with c3:
            st.markdown(f"**Duration:** {row['Duration'] or 'Verify'}")
        with c4:
            st.markdown(f"**Purpose:** {row['Purpose'] or 'Not specified'}")

        if row["Timing"]:
            st.markdown(f"**Timing:** {row['Timing']}")
        if row["Instructions"]:
            st.write(row["Instructions"])
        st.caption(f"Source: {row['Source']} | Confidence: {row['Confidence']}")
        st.divider()


def tab_protocols():
    st.header("Side-Effect Protocols & Urgent Rules")

    st.subheader("Side-Effect Protocols")
    protocols = get_protocols_df()
    render_df(protocols, "No protocols extracted yet.")

    if not protocols.empty:
        for _, row in protocols.iterrows():
            st.markdown(f"### {row['Protocol']}")
            st.markdown(f"**Trigger:** {row['Trigger']}")
            st.markdown("**Steps:**")
            st.markdown(row["Steps"] or "No steps extracted.")
            st.markdown(f"**Escalation:** {row['Escalation'] or 'Not specified'}")
            st.caption(f"Source: {row['Source']} | Confidence: {row['Confidence']}")
            st.divider()

    st.subheader("Urgent Red Flags")
    render_df(get_urgent_df(), "No urgent rules extracted yet.")


def tab_tasks():
    st.header("Tasks")

    df = get_tasks_df()
    render_df(df, "No care tasks extracted yet.")

    if not df.empty:
        st.subheader("Update Task Status")
        task_id = st.number_input("Task ID", min_value=1, step=1)
        new_status = st.selectbox("Status", ["pending", "done", "cancelled"])
        if st.button("Update Task"):
            execute("UPDATE care_tasks SET status = ? WHERE id = ?", (new_status, int(task_id)))
            st.success("Task updated.")
            st.rerun()


def tab_symptoms():
    st.header("Symptom Tracker")

    with st.form("symptom_form"):
        c1, c2, c3 = st.columns(3)
        with c1:
            symptom_date = st.date_input("Date", value=date.today())
        with c2:
            symptom_name = st.text_input("Symptom", placeholder="fever, loose stools, vomiting, pain")
        with c3:
            severity = st.slider("Severity", 1, 10, 3)

        temperature = st.text_input("Temperature, if relevant", placeholder="Example: 38.2C or 101F")
        notes = st.text_area("Notes", placeholder="Since when? How many episodes? Any weakness, bleeding, vomiting, etc.?")

        submitted = st.form_submit_button("Save Symptom")

    if submitted:
        if not symptom_name.strip():
            st.error("Enter a symptom.")
        else:
            flags = evaluate_symptom(symptom_name, severity, temperature, notes)
            save_symptom(symptom_date.isoformat(), symptom_name.strip(), severity, temperature.strip(), notes.strip(), flags)

            if flags:
                st.error("Urgent risk flags detected:")
                for f in flags:
                    st.markdown(f"- {f}")
            else:
                st.success("Symptom saved. No urgent rule-based flag detected.")

    st.subheader("Symptom History")
    render_df(get_symptoms_df(), "No symptoms tracked yet.")


def tab_memory():
    st.header("Caregiver Memory")

    st.write(
        "Use this for durable caregiver context: doctor instructions, patient preferences, food tolerance, side effects, "
        "hospital logistics, insurance notes, family decisions."
    )

    with st.form("memory_form"):
        memory_type = st.selectbox(
            "Memory Type",
            [
                "doctor_instruction",
                "patient_preference",
                "caregiver_observation",
                "medicine_note",
                "food_tolerance",
                "hospital_logistics",
                "insurance_note",
                "family_decision",
                "other",
            ]
        )
        title = st.text_input("Title")
        details = st.text_area("Details")
        submitted = st.form_submit_button("Save Memory")

    if submitted:
        if not title.strip() or not details.strip():
            st.error("Enter title and details.")
        else:
            execute("""
                INSERT INTO caregiver_memory
                (memory_type, title, details, created_at, updated_at)
                VALUES (?, ?, ?, ?, ?)
            """, (
                memory_type,
                title.strip(),
                details.strip(),
                now_iso(),
                now_iso()
            ))
            st.success("Memory saved.")

    st.subheader("Saved Memory")
    render_df(get_memory_df(), "No memory saved yet.")


def tab_brief():
    st.header("Doctor Visit / Caregiver Brief")

    if st.button("Generate Unified Case Brief", type="primary"):
        with st.spinner("Generating case brief..."):
            brief = generate_case_brief()
            st.markdown(brief)

    st.subheader("Export")

    export_json = export_case_json()
    st.download_button(
        "Download Full Case JSON",
        data=export_json,
        file_name=f"caregiver_case_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
        mime="application/json"
    )


def tab_raw():
    st.header("Raw Data")

    sections = {
        "Documents": get_documents_df(),
        "Case Profile": get_profile_df(),
        "Conflicts": get_conflicts_df(),
        "Diagnosis": get_diagnosis_df(),
        "Treatment Plans": get_treatment_df(),
        "Timeline": get_timeline_df(),
        "Scans": get_scans_df(),
        "Labs": get_labs_df(),
        "Medications": get_meds_df(),
        "Protocols": get_protocols_df(),
        "Urgent Rules": get_urgent_df(),
        "Contacts": get_contacts_df(),
        "Tasks": get_tasks_df(),
        "Symptoms": get_symptoms_df(),
        "Memory": get_memory_df(),
    }

    for name, df in sections.items():
        with st.expander(name, expanded=False):
            render_df(df, f"No {name.lower()} data.")


# ============================================================
# MAIN
# ============================================================

def main():
    init_db()

    st.set_page_config(
        page_title=APP_TITLE,
        page_icon="🩺",
        layout="wide"
    )

    st.title(APP_TITLE)
    warning_banner()

    with st.sidebar:
        st.header("System")

        if client:
            st.success(f"OpenAI connected: {OPENAI_MODEL}")
        else:
            st.error("OpenAI not configured. Add OPENAI_API_KEY to .env")

        st.divider()

        st.markdown("### Processing Model")
        st.markdown("""
        1. Upload document  
        2. Classify document type  
        3. Extract only relevant objects  
        4. Merge into unified patient case  
        5. Generate caregiver brief  
        """)

        st.divider()

        if st.button("Reset Database Completely"):
            reset_database()
            st.success("Database reset.")
            st.rerun()

    tabs = st.tabs([
        "Upload",
        "Dashboard",
        "Profile",
        "Timeline",
        "Scans",
        "Labs",
        "Medicines",
        "Protocols",
        "Tasks",
        "Symptoms",
        "Memory",
        "Brief",
        "Raw"
    ])

    with tabs[0]:
        tab_upload()

    with tabs[1]:
        tab_dashboard()

    with tabs[2]:
        tab_profile()

    with tabs[3]:
        tab_timeline()

    with tabs[4]:
        tab_scans()

    with tabs[5]:
        tab_labs()

    with tabs[6]:
        tab_medicines()

    with tabs[7]:
        tab_protocols()

    with tabs[8]:
        tab_tasks()

    with tabs[9]:
        tab_symptoms()

    with tabs[10]:
        tab_memory()

    with tabs[11]:
        tab_brief()

    with tabs[12]:
        tab_raw()


if __name__ == "__main__":
    main()